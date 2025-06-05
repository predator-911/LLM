import uuid
import logging
from typing import List, Dict, Tuple
import PyPDF2
import docx
from io import BytesIO
import re
from datetime import datetime

from config import Config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = Config.CHUNK_SIZE
        self.chunk_overlap = Config.CHUNK_OVERLAP
    
    async def process_document(self, content: bytes, filename: str) -> Tuple[str, List[Dict]]:
        """Process a document and return document ID and chunks"""
        document_id = str(uuid.uuid4())
        
        # Extract text based on file type
        text = await self._extract_text(content, filename)
        
        if not text.strip():
            raise ValueError("No text content found in document")
        
        # Split into chunks
        chunks = self._create_chunks(text, document_id, filename)
        
        logger.info(f"Processed {filename}: {len(chunks)} chunks created")
        return document_id, chunks
    
    async def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from different file types"""
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text(content)
            elif file_extension == 'docx':
                return self._extract_docx_text(content)
            elif file_extension in ['txt', 'md']:
                return content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ValueError(f"Failed to extract text from {filename}: {str(e)}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise ValueError(f"Failed to read PDF: {str(e)}")
        
        return text
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise ValueError(f"Failed to read DOCX: {str(e)}")
    
    def _create_chunks(self, text: str, document_id: str, filename: str) -> List[Dict]:
        """Split text into overlapping chunks"""
        chunks = []
        
        # Clean text
        text = self._clean_text(text)
        
        # Split into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        current_length = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed chunk size, save current chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunks.append(self._create_chunk_dict(
                    current_chunk.strip(),
                    document_id,
                    filename,
                    chunk_index
                ))
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + " " + sentence
                current_length = len(current_chunk)
                chunk_index += 1
            else:
                current_chunk += " " + sentence
                current_length += sentence_length
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append(self._create_chunk_dict(
                current_chunk.strip(),
                document_id,
                filename,
                chunk_index
            ))
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get the last N characters for overlap"""
        if len(text) <= overlap_size:
            return text
        return text[-overlap_size:]
    
    def _create_chunk_dict(self, content: str, document_id: str, filename: str, chunk_index: int) -> Dict:
        """Create a chunk dictionary"""
        return {
            'id': f"{document_id}_{chunk_index}",
            'content': content,
            'document_id': document_id,
            'filename': filename,
            'chunk_index': chunk_index,
            'length': len(content),
            'created_at': datetime.utcnow().isoformat()
        }
    
    async def get_page_count(self, content: bytes, filename: str) -> int:
        """Get the number of pages in a document"""
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                return len(pdf_reader.pages)
            elif file_extension == 'docx':
                # For DOCX, estimate based on content length
                doc = docx.Document(BytesIO(content))
                total_chars = sum(len(p.text) for p in doc.paragraphs)
                # Rough estimate: 2000 characters per page
                return max(1, total_chars // 2000)
            else:
                # For text files, estimate based on content length
                text = content.decode('utf-8')
                # Rough estimate: 2000 characters per page
                return max(1, len(text) // 2000)
        except Exception as e:
            logger.error(f"Error counting pages: {str(e)}")
            return 1  # Default to 1 page if we can't determine